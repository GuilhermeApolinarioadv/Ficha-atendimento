import streamlit as st
from docx import Document
from datetime import datetime
import io
import smtplib
from email.message import EmailMessage
import re

# Mapeamento de meses para português
MESES_PT = {
    "January": "janeiro", "February": "fevereiro", "March": "março",
    "April": "abril", "May": "maio", "June": "junho",
    "July": "julho", "August": "agosto", "September": "setembro",
    "October": "outubro", "November": "novembro", "December": "dezembro"
}

def formatar_data_portugues(data):
    dia = data.day
    mes = MESES_PT[data.strftime("%B")]
    ano = data.year
    return f"{dia} de {mes} de {ano}"

# Função para gerar o contrato preenchido
def gerar_contrato(dados):
    doc = Document("Contrato_Modelo_Final_Completo_Atualizado.docx")

    for p in doc.paragraphs:
        for chave, valor in dados.items():
            if chave in p.text:
                for run in p.runs:
                    if chave in run.text:
                        run.text = run.text.replace(chave, valor)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for chave, valor in dados.items():
                    if chave in cell.text:
                        cell.text = cell.text.replace(chave, valor)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Função para envio de e-mail
def enviar_email(arquivo, nome_cliente):
    email_de = "contratosguilherme.enviador@gmail.com"
    senha = "quprkjbbttuyxwnv"
    email_para = "contratosguilherme.enviador@gmail.com"

    msg = EmailMessage()
    msg["Subject"] = f"Ficha de Atendimento - {nome_cliente}"
    msg["From"] = email_de
    msg["To"] = email_para
    msg.set_content(f"Segue em anexo a ficha de atendimento gerada para {nome_cliente}.")

    msg.add_attachment(arquivo.getvalue(),
                       maintype='application',
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename="Ficha_Atendimento.docx")

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(email_de, senha)
        smtp.send_message(msg)

# Interface Streamlit
st.title("Formulário de Geração de Ficha de Atendimento")
st.markdown("Preencha as informações abaixo. A ficha de atendimento será enviada automaticamente por e-mail.")

dados = {}
dados["{{CONTRATANTE_NOME}}"] = st.text_input("Nome do Contratante")
cpf_input = st.text_input("CPF (formato: 000.000.000-00)")

if cpf_input and not re.match(r'^\d{3}\.\d{3}\.\d{3}-\d{2}$', cpf_input):
    st.error("CPF inválido. Use o formato 000.000.000-00")
dados["{{CPF}}"] = cpf_input

rg_input = st.text_input("RG (apenas números)", max_chars=20)
if rg_input and not rg_input.isdigit():
    st.error("RG deve conter apenas números.")
dados["{{RG}}"] = rg_input

dados["{{EMAIL}}"] = st.text_input("Email")
dados["{{NACIONALIDADE}}"] = st.text_input("Nacionalidade")
dados["{{ESTADO_CIVIL}}"] = st.text_input("Estado Civil")
dados["{{PROFISSAO}}"] = st.text_input("Profissão")
dados["{{RUA}}"] = st.text_input("Rua")
numero = st.text_input("Número (somente números)", max_chars=10)
if numero and not numero.isdigit():
    st.error("Número deve conter apenas dígitos.")
dados["{{NÚMERO}}"] = numero
dados["{{BAIRRO}}"] = st.text_input("Bairro")
dados["{{CIDADE}}"] = st.text_input("Cidade")
dados["{{ESTADO}}"] = st.text_input("Estado")
dados["{{CEP}}"] = st.text_input("CEP", max_chars=10)

data_assinatura = st.date_input("Data")
dados["{{DATA_ASSINATURA}}"] = formatar_data_portugues(data_assinatura)
dados["{{TABELA_PARCELAS}}"] = ""

if "enviado" not in st.session_state:
    st.session_state.enviado = False

if not st.session_state.enviado:
    if st.button("Enviar Formulário"):
        if re.match(r'^\d{3}\.\d{3}\.\d{3}-\d{2}$', cpf_input) and numero.isdigit() and rg_input.isdigit():
            contrato = gerar_contrato(dados)
            enviar_email(contrato, dados["{{CONTRATANTE_NOME}}"])
            st.success("Ficha de atendimento enviada com sucesso!")
            st.session_state.enviado = True
        else:
            st.warning("Corrija os campos inválidos antes de enviar.")
else:
    st.info("O formulário já foi enviado com sucesso.")
